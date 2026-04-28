from html import escape
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate


BASE = Path(__file__).resolve().parent
ROOT = BASE.parent
SOURCE = BASE / "Meidie_Fei_Cyber_Security_Resume.md"
PDF_TARGET = ROOT / "assets" / "Meidie_Fei_Cyber_Security_Resume.pdf"
DOCX_TARGET = BASE / "generated" / "Meidie_Fei_Cyber_Security_Resume.docx"


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((https?://.*?)\)", r"\1", text)
    return text.strip()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F5FBF")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_content(paragraph, text: str) -> None:
    text = text.replace("`", "")
    pattern = re.compile(r"(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*(.*?)\*\*)")
    position = 0

    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])

        if match.group(2) and match.group(3):
            add_hyperlink(paragraph, match.group(2), match.group(3))
        else:
            run = paragraph.add_run(match.group(4))
            run.bold = True

        position = match.end()

    if position < len(text):
        paragraph.add_run(text[position:])


def build_docx() -> None:
    DOCX_TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        raw = line.rstrip()
        if not raw:
            continue

        if raw.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_inline(raw[2:]))
            run.bold = True
            run.font.size = Pt(16)
            p.paragraph_format.space_after = Pt(2)
        elif raw.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(clean_inline(raw[3:]).upper())
            run.bold = True
            run.font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
        elif raw.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(1.5)
            add_inline_content(p, raw[2:].strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            stripped = raw.strip()
            if stripped.startswith("**") and "**" in stripped[2:]:
                match = re.match(r"\*\*(.*?)\*\*(.*)", stripped)
                if match:
                    run = p.add_run(clean_inline(match.group(1)))
                    run.bold = True
                    add_inline_content(p, match.group(2))
                else:
                    add_inline_content(p, stripped)
            else:
                if " | " in stripped or stripped.startswith("Email:") or stripped == "Melbourne, Australia":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_content(p, stripped)

    doc.save(DOCX_TARGET)


def inline_markup(text: str) -> str:
    text = text.replace("`", "")
    parts = re.split(r"(\[.*?\]\(https?://.*?\)|\*\*.*?\*\*)", text)
    rendered = []
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            rendered.append(f"<b>{escape(part[2:-2])}</b>")
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            match = re.match(r"\[(.*?)\]\((https?://.*?)\)", part)
            if match:
                rendered.append(f'<a href="{escape(match.group(2))}" color="#1f5fbf">{escape(match.group(1))}</a>')
            else:
                rendered.append(escape(part))
        else:
            rendered.append(escape(part))
    return "".join(rendered)


def styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=18,
            alignment=1,
            spaceAfter=2,
            textColor=colors.HexColor("#111827"),
        ),
        "normal": ParagraphStyle(
            "Normal",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.15,
            leading=9.4,
            spaceAfter=2.2,
            textColor=colors.HexColor("#111827"),
        ),
        "center": ParagraphStyle(
            "Center",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.15,
            leading=9.4,
            alignment=1,
            spaceAfter=2.2,
            textColor=colors.HexColor("#111827"),
        ),
        "heading": ParagraphStyle(
            "Heading",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=9.4,
            leading=10.4,
            spaceBefore=5,
            spaceAfter=2,
            textColor=colors.HexColor("#0f172a"),
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.15,
            leading=9.4,
            leftIndent=14,
            firstLineIndent=-8,
            bulletIndent=2,
            spaceAfter=1.2,
            textColor=colors.HexColor("#111827"),
        ),
    }


def build_pdf() -> None:
    PDF_TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(PDF_TARGET),
        pagesize=A4,
        rightMargin=0.45 * inch,
        leftMargin=0.45 * inch,
        topMargin=0.42 * inch,
        bottomMargin=0.42 * inch,
    )
    style_map = styles()
    story = []

    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        raw = line.rstrip()
        if not raw:
            continue

        if raw.startswith("# "):
            story.append(Paragraph(inline_markup(raw[2:]), style_map["title"]))
        elif raw.startswith("## "):
            story.append(Paragraph(escape(raw[3:].upper()), style_map["heading"]))
        elif raw.startswith("- "):
            story.append(Paragraph(inline_markup(raw[2:]), style_map["bullet"], bulletText="-"))
        else:
            style = style_map["center"] if (" | " in raw or raw.startswith("Email:") or raw == "Melbourne, Australia") else style_map["normal"]
            story.append(Paragraph(inline_markup(raw), style))

    doc.build(story)


def main() -> None:
    build_docx()
    build_pdf()
    print(f"Wrote {PDF_TARGET}")
    print(f"Wrote {DOCX_TARGET}")


if __name__ == "__main__":
    main()
